from __future__ import annotations

import io
import json
import subprocess
from pathlib import Path

import httpx
import pytest
from docx import Document
from openpyxl import Workbook

from ard_ossie.adapters.filesystem import RepositoryPaths
from ard_ossie.application.contracts import MutationRecord, WorkflowContext
from ard_ossie.application.finalize import FinalizeRequest, FinalizeService
from ard_ossie.application.intake import IssueAuthorizationService, IssueIntakeService
from ard_ossie.application.processing import ProcessingRequest, ProcessingService
from ard_ossie.application.release_detection import (
    ReleaseDetectionRequest,
    ReleaseDetectionService,
)
from ard_ossie.application.release_dispatch import (
    ReleaseDispatchRequest,
    ReleaseDispatchService,
)
from ard_ossie.application.release_publication import (
    ReleasePublicationRequest,
    ReleasePublicationService,
)
from ard_ossie.github_event import prepare_issue_event
from ard_ossie.pipeline import process_product
from ard_ossie.ports.git import ChangedPaths, CommitResult, GitConflict
from ard_ossie.ports.github import (
    PullRequestState,
    ReleaseAssetState,
    ReleaseState,
    RepositoryState,
)
from ard_ossie.registry import Registry

FIXTURES = Path("tests/fixtures/github")


def document_bytes() -> bytes:
    buffer = io.BytesIO()
    document = Document()
    document.add_heading("Order semantics", level=1)
    document.add_paragraph("An order is a confirmed customer purchase.")
    document.save(buffer)
    return buffer.getvalue()


def dictionary_bytes() -> bytes:
    buffer = io.BytesIO()
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Dictionary"
    sheet.append(
        [
            "platform",
            "catalog",
            "schema",
            "table",
            "column",
            "data_type",
            "nullable",
            "pk",
            "description",
        ]
    )
    sheet.append(
        [
            "erp",
            "analytics",
            "sales",
            "orders",
            "order_id",
            "INT64",
            "false",
            "true",
            "Unique order identifier",
        ]
    )
    workbook.save(buffer)
    workbook.close()
    return buffer.getvalue()


def attachment_transport(expected_token: str) -> httpx.MockTransport:
    payloads = {
        "11111111-1111-1111-1111-111111111111": (
            b"<html><body><h1>Sales Order</h1><p>Order analytics.</p></body></html>",
            "text/html",
        ),
        "22222222-2222-2222-2222-222222222222": (
            document_bytes(),
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ),
        "33333333-3333-3333-3333-333333333333": (
            dictionary_bytes(),
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        ),
    }

    def handler(request: httpx.Request) -> httpx.Response:
        assert request.headers["authorization"] == f"Bearer {expected_token}"
        key = request.url.path.rsplit("/", 1)[-1]
        content, content_type = payloads[key]
        return httpx.Response(200, headers={"content-type": content_type}, content=content)

    return httpx.MockTransport(handler)


def git(repository: Path, *arguments: str) -> str:
    result = subprocess.run(
        ["git", *arguments],
        cwd=repository,
        check=True,
        capture_output=True,
        text=True,
    )
    return result.stdout.strip()


class LifecycleGit:
    def __init__(self, repository: Path) -> None:
        self.repository = repository
        self.remote_heads: dict[str, str] = {}
        self.pushes: list[tuple[str, bool]] = []
        self.pushed_tags: list[tuple[str, ...]] = []

    def switch_or_create(self, branch: str, base_ref: str) -> None:
        git(self.repository, "switch", "--create", branch, base_ref)

    def current_sha(self) -> str:
        return git(self.repository, "rev-parse", "HEAD")

    def remote_branch_sha(self, branch: str) -> str | None:
        return self.remote_heads.get(branch)

    def read_text_at(self, revision: str, path: str | Path) -> str:
        result = subprocess.run(
            ["git", "show", f"{revision}:{Path(path).as_posix()}"],
            cwd=self.repository,
            capture_output=True,
            text=True,
        )
        if result.returncode != 0:
            raise GitConflict("REVISION_FILE_NOT_FOUND", result.stderr)
        return result.stdout

    def read_bytes_at(self, revision: str, path: str | Path) -> bytes:
        result = subprocess.run(
            ["git", "show", f"{revision}:{Path(path).as_posix()}"],
            cwd=self.repository,
            capture_output=True,
        )
        if result.returncode != 0:
            raise GitConflict(
                "REVISION_FILE_NOT_FOUND",
                result.stderr.decode("utf-8", errors="replace"),
            )
        return result.stdout

    def commit_intake_paths(self, product_key: str, message: str) -> CommitResult:
        return self._commit(message, f"products/{product_key}")

    def commit_allowed_paths(self, product_key: str, message: str) -> CommitResult:
        return self._commit(
            message,
            f"products/{product_key}/generated",
            f"products/{product_key}/quality",
            "registry",
        )

    def push(self, branch: str, *, lfs: bool = False) -> None:
        self.remote_heads[branch] = self.current_sha()
        self.pushes.append((branch, lfs))

    def changed_paths(self, base_ref: str, head_ref: str = "HEAD") -> ChangedPaths:
        merge_base = git(self.repository, "merge-base", base_ref, head_ref)
        changed = git(
            self.repository,
            "diff",
            "--name-only",
            f"{merge_base}...{head_ref}",
        )
        return ChangedPaths(
            merge_base=merge_base,
            paths=tuple(Path(item) for item in changed.splitlines() if item),
        )

    def tag_target(self, tag: str) -> str | None:
        result = subprocess.run(
            ["git", "rev-parse", "--verify", "--quiet", f"refs/tags/{tag}^{{commit}}"],
            cwd=self.repository,
            capture_output=True,
            text=True,
        )
        return result.stdout.strip() if result.returncode == 0 else None

    def create_annotated_tag(self, tag: str, target: str, message: str) -> None:
        git(self.repository, "tag", "--annotate", tag, target, "--message", message)

    def push_tags(self, tags: list[str]) -> None:
        self.pushed_tags.append(tuple(tags))

    def is_ancestor(self, ancestor: str, descendant: str) -> bool:
        result = subprocess.run(
            ["git", "merge-base", "--is-ancestor", ancestor, descendant],
            cwd=self.repository,
        )
        return result.returncode == 0

    def _commit(self, message: str, *paths: str) -> CommitResult:
        git(self.repository, "add", "--", *paths)
        git(self.repository, "commit", "-m", message)
        return CommitResult(sha=self.current_sha(), created=True)


class LifecycleGitHub:
    def __init__(self, git_adapter: LifecycleGit) -> None:
        self.git = git_adapter
        self.labels = {"ard:approved"}
        self.pull_request: PullRequestState | None = None
        self.statuses: dict[tuple[str, str], str] = {}
        self.comments: dict[str, str] = {}
        self.releases: dict[str, ReleaseState] = {}
        self.dispatches: list[tuple[str, dict[str, object]]] = []

    def repository(self) -> RepositoryState:
        return RepositoryState(
            full_name="owner/repository",
            public=True,
            archived=False,
            default_branch="main",
            permission="admin",
        )

    def collaborator_permission(self, login: str) -> str:
        assert login == "maintainer"
        return "write"

    def find_open_pr(self, branch: str) -> PullRequestState | None:
        if self.pull_request is None or self.pull_request.head_branch != branch:
            return None
        return self._current_pr()

    def create_draft_pr(
        self,
        branch: str,
        base: str,
        title: str,
        body: str,
    ) -> PullRequestState:
        self.pull_request = PullRequestState(
            number=7,
            head_branch=branch,
            head_sha=self.git.remote_heads[branch],
            base_branch=base,
            draft=True,
            merged_at=None,
            merge_sha=None,
            url="https://example.invalid/pull/7",
        )
        return self.pull_request

    def get_pr(self, number: int) -> PullRequestState:
        assert number == 7
        return self._current_pr()

    def set_issue_labels(
        self,
        number: int,
        *,
        add: set[str],
        remove: set[str],
    ) -> list[MutationRecord]:
        assert number == 42
        before = set(self.labels)
        self.labels.difference_update(remove)
        self.labels.update(add)
        if before == self.labels:
            return []
        return [MutationRecord(resource="label", target="issue:42", action="set")]

    def set_status(
        self,
        sha: str,
        context: str,
        state: str,
        description: str,
        target_url: str,
    ) -> MutationRecord:
        self.statuses[(sha, context)] = state
        return MutationRecord(resource="status", target=f"{sha}:{context}", action="set")

    def get_status(self, sha: str, context: str) -> str | None:
        return self.statuses.get((sha, context))

    def dispatch_workflow(
        self,
        workflow: str,
        ref: str,
        inputs: dict[str, str],
    ) -> MutationRecord:
        raise AssertionError("this fixture must not require a changeset dispatch")

    def get_release(self, tag: str) -> ReleaseState | None:
        return self.releases.get(tag)

    def upsert_release(
        self,
        tag: str,
        title: str,
        asset: Path,
        sha256: str,
    ) -> MutationRecord:
        action = "noop" if tag in self.releases else "upload"
        self.releases[tag] = ReleaseState(
            id=1,
            tag=tag,
            title=title,
            draft=False,
            prerelease=False,
            assets=(
                ReleaseAssetState(
                    name=asset.name,
                    digest=sha256,
                    url="https://example.invalid/release.zip",
                ),
            ),
        )
        return MutationRecord(resource="release", target=tag, action=action)

    def repository_dispatch(
        self,
        event_type: str,
        payload: dict[str, object],
    ) -> MutationRecord:
        self.dispatches.append((event_type, payload))
        return MutationRecord(
            resource="repository_dispatch",
            target=event_type,
            action="dispatch",
        )

    def upsert_pr_comment(self, number: int, marker: str, body: str) -> MutationRecord:
        previous = self.comments.get(marker)
        self.comments[marker] = body
        action = "create" if previous is None else ("noop" if previous == body else "update")
        return MutationRecord(
            resource="comment",
            target=f"pr:{number}:{marker}",
            action=action,
        )

    def _current_pr(self) -> PullRequestState:
        assert self.pull_request is not None
        return PullRequestState(
            number=self.pull_request.number,
            head_branch=self.pull_request.head_branch,
            head_sha=self.git.remote_heads[self.pull_request.head_branch],
            base_branch=self.pull_request.base_branch,
            draft=self.pull_request.draft,
            merged_at=self.pull_request.merged_at,
            merge_sha=self.pull_request.merge_sha,
            url=self.pull_request.url,
        )


class RecordingProvider:
    def __init__(self) -> None:
        self.requests: list[dict[str, object]] = []

    def capabilities(self) -> dict[str, str]:
        return {
            "structured_output": "json_schema",
            "provider": "test-provider",
            "model": "test-model",
        }

    def generate_structured(self, *, schema, messages):
        self.requests.append({"schema": schema, "messages": messages})
        return {"suggestions": [], "metrics": [], "product_facts": []}


def test_approved_public_issue_with_auth_releases_reproducibly_and_traceably(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    repository = tmp_path / "repository"
    repository.mkdir()
    git(repository, "init", "-b", "main")
    git(repository, "config", "user.name", "ARD Test")
    git(repository, "config", "user.email", "ard-test@example.invalid")
    git(repository, "commit", "--allow-empty", "-m", "chore: initialize repository")
    base_head = git(repository, "rev-parse", "HEAD")

    event_payload = json.loads((FIXTURES / "approved-issue.json").read_text(encoding="utf-8"))
    event_payload["repository"] = {
        "default_branch": "main",
        "full_name": "owner/repository",
    }
    event_path = tmp_path / "approved-issue-event.json"
    event_path.write_text(json.dumps(event_payload), encoding="utf-8")
    context = WorkflowContext(
        repository=repository,
        event_path=event_path,
        event_name="issues",
        repository_name="owner/repository",
        actor="maintainer",
        runner_temp=tmp_path,
    )
    git_adapter = LifecycleGit(repository)
    git_adapter.remote_heads["main"] = base_head
    github = LifecycleGitHub(git_adapter)
    authorization = IssueAuthorizationService(github).run(
        context,
        label="ard:approved",
        actor="maintainer",
    )
    assert authorization.outputs == {"allowed": True}

    attachment_token = "e2e-attachment-token"
    monkeypatch.setenv("ARD_ATTACHMENT_TOKEN", attachment_token)
    with httpx.Client(transport=attachment_transport(attachment_token)) as client:
        intake_result = IssueIntakeService(
            RepositoryPaths(repository),
            git_adapter,
            github,
            prepare=lambda event, workspace: prepare_issue_event(
                event,
                workspace,
                client=client,
            ),
        ).run(context)
    product_key = str(intake_result.outputs["product_key"])
    product_id = str(intake_result.outputs["product_id"])
    branch = str(intake_result.outputs["branch"])
    intake_head = str(intake_result.outputs["expected_head"])
    pr_number = int(intake_result.outputs["pr_number"])
    product_root = repository / "products" / product_key

    provider = RecordingProvider()
    processing_result = ProcessingService(
        RepositoryPaths(repository),
        git_adapter,
        github,
        provider_factory=lambda: provider,
    ).run(
        ProcessingRequest(
            repository=repository,
            product_key=product_key,
            branch=branch,
            pr_number=pr_number,
            expected_head=intake_head,
            allow_writeback=True,
            target_url="https://example.invalid/actions/runs/1",
        )
    )
    processed_head = str(processing_result.outputs["current_head"])
    FinalizeService(RepositoryPaths(repository), github).run(
        FinalizeRequest(
            repository=repository,
            upstream_result="success",
            issue_number=42,
            pr_number=pr_number,
            expected_head=processed_head,
        )
    )

    first_dictionary = json.loads(
        (product_root / "generated" / "data-dictionary.json").read_text(encoding="utf-8")
    )
    first_column_id = first_dictionary["tables"][0]["columns"][0]["column_id"]
    first_generated = {
        path.name: path.read_bytes() for path in (product_root / "generated").iterdir()
    }
    second = process_product(product_root, registry_root=repository / "registry")
    second_dictionary = json.loads(
        (product_root / "generated" / "data-dictionary.json").read_text(encoding="utf-8")
    )
    assert second_dictionary["tables"][0]["columns"][0]["column_id"] == first_column_id
    assert {
        path.name: path.read_bytes() for path in (product_root / "generated").iterdir()
    } == first_generated

    registry = Registry.load(repository / "registry")
    product = registry.get_product(product_id)
    assert product is not None and product.version == 1
    tables = registry.tables()
    assert len(tables) == 1 and tables[0].version == 1
    assert registry.mappings()[0].product_id == product_id
    assert registry.mappings()[0].table_id == tables[0].table_id
    assert processing_result.outputs["product_id"] == second.product_id == product_id
    assert processing_result.outputs["version"] == 1
    assert github.pull_request is not None and github.pull_request.draft
    assert github.labels == {"ard:approved", "ard:pr-created"}
    assert github.statuses == {
        (processed_head, "ard/changeset"): "success",
        (processed_head, "ard/quality-gate"): "success",
    }
    assert {path.name for path in product_root.joinpath("generated").iterdir()} == {
        "data-product.md",
        "data-semantic.md",
        "data-dictionary.json",
        "ossie-model.json",
        "source-manifest.json",
    }
    assert {path.name for path in product_root.joinpath("quality").iterdir()} == {
        "quality-report.json",
        "duplicate-report.json",
        "version-report.json",
        "impact-report.json",
        "llm-suggestions.json",
        "semantic-fidelity.json",
    }
    assert provider.requests
    assert all(
        "semantic-structure-repair-v1" not in json.dumps(request) for request in provider.requests
    )

    manifest = json.loads(
        (product_root / "generated" / "source-manifest.json").read_text(encoding="utf-8")
    )
    assert len(manifest["files"]) == 3
    assert all(len(item["sha256"]) == 64 for item in manifest["files"])
    assert all(not item["relative_path"].startswith("/") for item in manifest["files"])

    detection = ReleaseDetectionService(
        RepositoryPaths(repository),
        git_adapter,
    ).run(
        ReleaseDetectionRequest(
            repository=repository,
            before=base_head,
            current=processed_head,
        )
    )
    assert detection.outputs["products"] == [product_key]
    assert detection.outputs["tables"] == [tables[0].table_id]
    publication = ReleasePublicationService(
        RepositoryPaths(repository),
        git_adapter,
        github,
    ).run(
        ReleasePublicationRequest(
            repository=repository,
            product_key=product_key,
            current=processed_head,
            table_ids=[tables[0].table_id],
            output=repository / "dist",
        )
    )
    assert publication.outputs["product_tag"] == f"product/{product_id}/v1"
    assert publication.outputs["table_tags"] == [f"table/{tables[0].table_id}/v1"]
    release_result = repository / ".ard" / "run" / "workflow.release-product-result.json"
    release_result.parent.mkdir(parents=True, exist_ok=True)
    release_result.write_text(publication.model_dump_json(), encoding="utf-8")
    dispatch = ReleaseDispatchService(
        RepositoryPaths(repository),
        git_adapter,
        github,
    ).run(
        ReleaseDispatchRequest(
            repository=repository,
            result_path=release_result,
            current=processed_head,
        )
    )
    assert dispatch.outputs["dispatched"] is True
    assert github.dispatches[0][0] == "ard_product_released"
    assert github.dispatches[0][1]["product_id"] == product_id
    assert git(repository, "rev-list", "--count", "HEAD") == "3"
    assert git(repository, "tag", "--list", "*/v1").splitlines() == [
        publication.outputs["product_tag"],
        publication.outputs["table_tags"][0],
    ]
    for relative_path in git(repository, "ls-files").splitlines():
        secret_assignment = b"ARD_LLM_API_KEY" + b"="
        assert secret_assignment not in (repository / relative_path).read_bytes()
