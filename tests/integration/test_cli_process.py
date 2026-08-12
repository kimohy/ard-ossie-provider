from __future__ import annotations

import json
from importlib import import_module
from pathlib import Path

import pytest
import yaml
from docx import Document
from jsonschema import validate
from openpyxl import Workbook
from typer.testing import CliRunner

from ard_ossie.cli import app
from ard_ossie.ossie_compiler import load_ossie_011_schema
from ard_ossie.pipeline import (
    ProviderExecutionError,
    ProviderFailureKind,
    process_product,
)

PRODUCT_ID = "prd_0198f6c2-8ac7-7f31-a48e-1c3d82e9a631"
TABLE_ID = "tbl_0198f6ca-2a11-78d1-8672-67d49e69f14c"
CUSTOMERS_TABLE_ID = "tbl_0198f6ca-2a11-78d1-8672-67d49e69f14d"
process_cli = import_module("ard_ossie.cli.process")


def create_product_fixture(root: Path, *, valid_dictionary: bool = True) -> Path:
    product = root / "products" / "sales-order"
    for directory in ("product-info", "semantic", "dictionary"):
        (product / "sources" / directory).mkdir(parents=True, exist_ok=True)
    (product / "sources" / "product-info" / "product.html").write_text(
        "<html><body><h1>Sales Order</h1><p>Order analytics.</p></body></html>",
        encoding="utf-8",
    )
    document = Document()
    document.add_heading("Order semantics", level=1)
    document.add_paragraph("An order is a confirmed customer purchase.")
    document.save(product / "sources" / "semantic" / "semantic.docx")

    workbook = Workbook()
    sheet = workbook.active
    headers = ["platform", "catalog", "schema", "table", "column", "data_type", "nullable"]
    if valid_dictionary:
        headers.append("pk")
    sheet.append(headers)
    sheet.append(["erp", "analytics", "sales", "orders", "order_id", "INT64", "false", "true"])
    workbook.save(product / "sources" / "dictionary" / "dictionary.xlsx")

    config = {
        "operation": "create",
        "product_id": PRODUCT_ID,
        "product_key": "sales-order",
        "version": 1,
        "display_name": "Sales Order",
        "description": "Order analytics product.",
        "tables": [
            {
                "locator": "erp|analytics|sales|orders",
                "table_id": TABLE_ID,
                "version": 1,
                "usage": "SOURCE",
            }
        ],
    }
    (product / "product.yaml").write_text(yaml.safe_dump(config, sort_keys=False), encoding="utf-8")
    return product


def test_process_emits_required_artifacts_and_reuses_column_ids(tmp_path: Path) -> None:
    product = create_product_fixture(tmp_path)
    registry = tmp_path / "registry"
    runner = CliRunner()

    first = runner.invoke(app, ["process", str(product), "--registry", str(registry)])

    assert first.exit_code == 0, first.output
    generated = product / "generated"
    assert {path.name for path in generated.iterdir()} == {
        "data-product.md",
        "data-semantic.md",
        "data-dictionary.json",
        "ossie-model.json",
        "source-manifest.json",
    }
    first_dictionary = json.loads((generated / "data-dictionary.json").read_text())
    column_id = first_dictionary["tables"][0]["columns"][0]["column_id"]
    assert "confirmed customer purchase" in (generated / "data-semantic.md").read_text()
    assert "Order analytics" in (generated / "data-product.md").read_text()
    first_artifacts = {path.name: path.read_bytes() for path in generated.iterdir()}

    second = runner.invoke(app, ["process", str(product), "--registry", str(registry)])
    second_dictionary = json.loads((generated / "data-dictionary.json").read_text())

    assert second.exit_code == 0, second.output
    assert {path.name: path.read_bytes() for path in generated.iterdir()} == first_artifacts
    assert second_dictionary["tables"][0]["columns"][0]["column_id"] == column_id
    table_record = json.loads((registry / "tables" / f"{TABLE_ID}.json").read_text())
    assert table_record["columns"][0]["column_id"] == column_id
    quality = json.loads((product / "quality" / "quality-report.json").read_text())
    assert quality["status"] in {"PASS", "WARN"}


def test_process_cli_passes_warnings_as_errors_before_promotion(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    product = create_product_fixture(tmp_path)
    registry = tmp_path / "registry"
    captured: dict[str, object] = {}
    real_process_product = process_cli.process_product

    def capture_process_product(*args, **kwargs):
        captured.update(kwargs)
        return real_process_product(*args, **kwargs)

    monkeypatch.setattr(process_cli, "process_product", capture_process_product)
    monkeypatch.setattr(process_cli, "_provider_from_environment", lambda: None)

    result = CliRunner().invoke(
        app,
        [
            "process",
            str(product),
            "--registry",
            str(registry),
            "--warnings-as-errors",
        ],
    )

    assert captured["warnings_as_errors"] is True
    assert result.exit_code == 2
    assert not (product / "generated").exists()
    assert not (registry / "products").exists()


class FakeSemanticProvider:
    def health_check(self) -> bool:
        return True

    def capabilities(self) -> dict[str, str]:
        return {"structured_output": "json_schema"}

    def generate_structured(self, *, schema, messages):
        source = json.loads(messages[1]["content"])["semantic"]
        allowed_paths = json.loads(
            messages[0]["content"].split("Allowed field_path values: ", maxsplit=1)[1]
        )
        column_description_path = next(
            path
            for path in allowed_paths
            if path.startswith(f"tables.{TABLE_ID}.columns.")
            and path.endswith(".description")
        )
        evidence = [
            {
                "source_hash": source["source_hash"],
                "role": "semantic_document",
                "locator": {"document": "semantic/semantic.docx"},
                "excerpt": "confirmed customer purchase",
            }
        ]
        return {
            "suggestions": [
                {
                    "field_path": "product.synonyms",
                    "value": ["purchase orders"],
                    "confidence": 0.95,
                    "evidence": evidence,
                    "status": "ai_suggested",
                },
                {
                    "field_path": f"tables.{TABLE_ID}.description",
                    "value": "Confirmed orders",
                    "confidence": 0.9,
                    "evidence": evidence,
                    "status": "ai_suggested",
                },
                {
                    "field_path": column_description_path,
                    "value": "LLM-only order identifier",
                    "confidence": 0.9,
                    "evidence": evidence,
                    "status": "ai_suggested",
                },
            ],
            "metrics": [],
            "product_facts": [],
        }


class NoisyPortalProductProvider(FakeSemanticProvider):
    def generate_structured(self, *, schema, messages):
        product = json.loads(messages[1]["content"])["product"]

        def evidence_id_for(text: str) -> str:
            item = next(
                item
                for item in product["evidence"]
                if text in str(item.get("excerpt") or "")
            )
            return str(item["evidence_id"])

        return {
            "suggestions": [],
            "metrics": [],
            "product_facts": [
                {
                    "kind": "purpose",
                    "value": "주문 분석 지원",
                    "confidence": 0.98,
                    "evidence_ids": [evidence_id_for("주문 분석 지원")],
                },
                {
                    "kind": "domain",
                    "value": "영업",
                    "confidence": 0.97,
                    "evidence_ids": [evidence_id_for("도메인: 영업")],
                },
                {
                    "kind": "tag",
                    "value": "주문",
                    "confidence": 0.96,
                    "evidence_ids": [evidence_id_for("태그: 주문")],
                },
            ],
        }


def test_pipeline_normalizes_user_facts_and_excludes_portal_boilerplate(
    tmp_path: Path,
) -> None:
    product = create_product_fixture(tmp_path)
    portal_sentinels = (
        "데이터 상품 홈",
        "작성 도움말을 확인하세요",
        "첨부파일 다운로드 12KB",
        "개인정보를 입력하지 마세요",
        "빈 입력 필드",
        "AI가 만든 원문에 없는 요약",
        "이전 다음",
        "포털 푸터",
        "무엇을 도와드릴까요",
    )
    (product / "sources" / "product-info" / "product.html").write_text(
        """<html><body>
          <nav>데이터 상품 홈</nav>
          <main>
            <h1>Sales Order</h1>
            <p>목적: 주문 분석 지원</p>
            <p>도메인: 영업</p>
            <p>태그: 주문</p>
            <p>작성 도움말을 확인하세요</p>
            <button>첨부파일 다운로드 12KB</button>
            <aside>개인정보를 입력하지 마세요</aside>
            <label>빈 입력 필드</label><input value="">
            <p>(AI 자동생성) AI가 만든 원문에 없는 요약</p>
          </main>
          <a>이전 다음</a><footer>포털 푸터</footer>
          <section>무엇을 도와드릴까요</section>
        </body></html>""",
        encoding="utf-8",
    )

    process_product(
        product,
        registry_root=tmp_path / "registry",
        provider=NoisyPortalProductProvider(),
    )

    generated = product / "generated"
    markdown = (generated / "data-product.md").read_text(encoding="utf-8")
    assert "**Purpose:** 주문 분석 지원" in markdown
    assert "**Domain:** 영업" in markdown
    assert "**Tag:** 주문" in markdown
    assert "## Parsed source" not in markdown
    for sentinel in portal_sentinels:
        assert sentinel not in markdown

    ossie = json.loads((generated / "ossie-model.json").read_text(encoding="utf-8"))
    validate(instance=ossie, schema=load_ossie_011_schema())
    audit = json.loads(
        (product / "quality" / "llm-suggestions.json").read_text(encoding="utf-8")
    )
    assert [fact["kind"] for fact in audit["product_facts"]] == [
        "purpose",
        "domain",
        "tag",
    ]
    assert audit["product_facts"][0]["confidence"] == 0.98
    assert audit["product_facts"][0]["evidence"][0]["excerpt"]
    assert all("evidence_ids" not in fact for fact in audit["product_facts"])


def test_pipeline_audits_table_and_column_suggestions_without_publishing_them(
    tmp_path: Path,
) -> None:
    product = create_product_fixture(tmp_path)

    process_product(
        product,
        registry_root=tmp_path / "registry",
        provider=FakeSemanticProvider(),
    )

    ossie = json.loads((product / "generated" / "ossie-model.json").read_text())
    model = ossie["semantic_model"][0]
    assert model["ai_context"]["synonyms"] == ["purchase orders"]
    assert "description" not in model["datasets"][0]

    dictionary = json.loads(
        (product / "generated" / "data-dictionary.json").read_text(encoding="utf-8")
    )
    assert dictionary["tables"][0]["description"] is None
    assert dictionary["tables"][0]["columns"][0]["description"] is None

    audit = json.loads(
        (product / "quality" / "llm-suggestions.json").read_text(encoding="utf-8")
    )
    assert {
        item["value"]
        for item in audit["suggestions"]
        if isinstance(item["value"], str)
    } >= {
        "Confirmed orders",
        "LLM-only order identifier",
    }


class PhysicalFieldProvider(FakeSemanticProvider):
    def generate_structured(self, *, schema, messages):
        response = super().generate_structured(schema=schema, messages=messages)
        response["suggestions"][0]["field_path"] = (
            f"tables.{TABLE_ID}.columns.order_id.data_type"
        )
        return response


def test_pipeline_classifies_semantic_output_validation_without_leaking_value(
    tmp_path: Path,
) -> None:
    product = create_product_fixture(tmp_path)

    with pytest.raises(ProviderExecutionError) as captured:
        process_product(
            product,
            registry_root=tmp_path / "registry",
            provider=PhysicalFieldProvider(),
        )

    assert captured.value.code == "LLM_PHYSICAL_FIELD_FORBIDDEN"
    assert captured.value.kind is ProviderFailureKind.OUTPUT
    assert "data_type" not in str(captured.value)


class FakeMetricProvider:
    def health_check(self) -> bool:
        return True

    def capabilities(self) -> dict[str, str]:
        return {"structured_output": "json_schema"}

    def generate_structured(self, *, schema, messages):
        source = json.loads(messages[1]["content"])["semantic"]
        evidence = [
            {
                "source_hash": source["source_hash"],
                "role": "semantic_document",
                "locator": {"document": "semantic/semantic.docx"},
                "excerpt": "Order count is the number of orders",
            }
        ]
        return {
            "suggestions": [],
            "metrics": [
                {
                    "name": "order_count",
                    "expression": "COUNT(orders.order_id)",
                    "dataset_names": ["orders"],
                    "description": "Number of orders",
                    "synonyms": ["orders"],
                    "confidence": 0.96,
                    "evidence": evidence,
                    "status": "ai_suggested",
                }
            ],
            "product_facts": [],
        }


class DatasetSafetyProvider:
    def health_check(self) -> bool:
        return True

    def capabilities(self) -> dict[str, str]:
        return {"structured_output": "json_schema"}

    def generate_structured(self, *, schema, messages):
        source = json.loads(messages[1]["content"])["semantic"]
        evidence = [
            {
                "source_hash": source["source_hash"],
                "role": "semantic_document",
                "locator": {"document": "semantic/semantic.docx"},
                "excerpt": "Campaign count and modeled efficiency definitions",
            }
        ]
        return {
            "suggestions": [],
            "metrics": [
                {
                    "name": "Campaign Count",
                    "expression": "COUNT(DISTINCT campaign_id)",
                    "dataset_names": ["marketing_campaign"],
                    "description": "Distinct campaigns",
                    "synonyms": [],
                    "confidence": 0.98,
                    "evidence": evidence,
                    "status": "ai_suggested",
                },
                {
                    "name": "Modeled Efficiency",
                    "expression": (
                        "SUM(marketing_campaign.revenue) / SUM(sales_order.cost)"
                    ),
                    "dataset_names": ["marketing_campaign", "sales_order"],
                    "description": "Revenue divided by sales cost",
                    "synonyms": [],
                    "confidence": 0.95,
                    "evidence": evidence,
                    "status": "ai_suggested",
                },
            ],
            "product_facts": [],
        }


class UnsafeMetricProvider(DatasetSafetyProvider):
    def generate_structured(self, *, schema, messages):
        response = super().generate_structured(schema=schema, messages=messages)
        response["metrics"][0]["expression"] = "DELETE FROM marketing_campaign"
        response["metrics"][0]["confidence"] = 0.1
        return response


class DuplicateMetricProvider(DatasetSafetyProvider):
    def generate_structured(self, *, schema, messages):
        response = super().generate_structured(schema=schema, messages=messages)
        response["metrics"].append(dict(response["metrics"][0]))
        return response


class InitiallyAcceptedMetricsProvider(DatasetSafetyProvider):
    def generate_structured(self, *, schema, messages):
        response = super().generate_structured(schema=schema, messages=messages)
        response["metrics"][1]["expression"] = "SUM(marketing_campaign.revenue)"
        response["metrics"][1]["dataset_names"] = ["marketing_campaign"]
        return response


def configure_metric_safety_fixture(product: Path) -> None:
    workbook = Workbook()
    sheet = workbook.active
    sheet.append(
        [
            "platform",
            "catalog",
            "schema",
            "table",
            "column",
            "data_type",
            "nullable",
            "pk",
            "description",
        ]
    )
    sheet.append(
        [
            "warehouse",
            "analytics",
            "marketing",
            "marketing_campaign",
            "campaign_id",
            "STRING",
            False,
            True,
            "Campaign identifier",
        ]
    )
    sheet.append(
        [
            "warehouse",
            "analytics",
            "marketing",
            "marketing_campaign",
            "revenue",
            "NUMERIC",
            False,
            False,
            "Attributed revenue",
        ]
    )
    sheet.append(
        [
            "warehouse",
            "analytics",
            "sales",
            "sales_order",
            "order_id",
            "STRING",
            False,
            True,
            "Order identifier",
        ]
    )
    sheet.append(
        [
            "warehouse",
            "analytics",
            "sales",
            "sales_order",
            "campaign_id",
            "STRING",
            False,
            False,
            "Campaign identifier",
        ]
    )
    sheet.append(
        [
            "warehouse",
            "analytics",
            "sales",
            "sales_order",
            "cost",
            "NUMERIC",
            False,
            False,
            "Sales cost",
        ]
    )
    workbook.save(product / "sources" / "dictionary" / "dictionary.xlsx")
    config = yaml.safe_load((product / "product.yaml").read_text(encoding="utf-8"))
    config["tables"] = [
        {
            "locator": "warehouse|analytics|marketing|marketing_campaign",
            "table_id": CUSTOMERS_TABLE_ID,
            "version": 1,
            "usage": "SOURCE",
        },
        {
            "locator": "warehouse|analytics|sales|sales_order",
            "table_id": TABLE_ID,
            "version": 1,
            "usage": "SOURCE",
        },
    ]
    (product / "product.yaml").write_text(
        yaml.safe_dump(config, sort_keys=False),
        encoding="utf-8",
    )


def test_pipeline_qualifies_single_dataset_metric_and_excludes_cross_dataset_metric(
    tmp_path: Path,
) -> None:
    product = create_product_fixture(tmp_path)
    configure_metric_safety_fixture(product)
    registry = tmp_path / "registry"

    result = process_product(
        product,
        registry_root=registry,
        provider=DatasetSafetyProvider(),
    )

    ossie = json.loads((product / "generated" / "ossie-model.json").read_text())
    model = ossie["semantic_model"][0]
    published_metrics = [
        (
            metric["name"],
            metric["expression"]["dialects"][0]["expression"],
        )
        for metric in model["metrics"]
    ]
    assert published_metrics == [
        (
            "Campaign Count",
            "COUNT(DISTINCT marketing_campaign.campaign_id)",
        )
    ]
    assert model["relationships"] == []
    semantic_markdown = (product / "generated" / "data-semantic.md").read_text()
    assert "confirmed customer purchase" in semantic_markdown
    assert "Campaign Count" not in semantic_markdown
    assert "Modeled Efficiency" not in semantic_markdown

    product_record = json.loads(
        (registry / "products" / f"{PRODUCT_ID}.json").read_text()
    )
    assert [metric["name"] for metric in product_record["metrics"]] == [
        "Campaign Count"
    ]
    metric_findings = [
        finding
        for finding in result.quality_report.warnings
        if finding.code == "METRIC_MULTI_DATASET_UNSUPPORTED"
    ]
    assert result.quality_report.status == "WARN"
    assert len(metric_findings) == 1
    assert metric_findings[0].path == "metrics.Modeled Efficiency"

    audit = json.loads(
        (product / "quality" / "llm-suggestions.json").read_text(encoding="utf-8")
    )
    assert [metric["name"] for metric in audit["metrics"]] == [
        "Campaign Count",
        "Modeled Efficiency",
    ]
    assert [metric["expression"] for metric in audit["metrics"]] == [
        "COUNT(DISTINCT campaign_id)",
        "SUM(marketing_campaign.revenue) / SUM(sales_order.cost)",
    ]
    assert [metric["dataset_names"] for metric in audit["metrics"]] == [
        ["marketing_campaign"],
        ["marketing_campaign", "sales_order"],
    ]


def test_pipeline_rejects_unsafe_metric_before_confidence_filter_or_promotion(
    tmp_path: Path,
) -> None:
    product = create_product_fixture(tmp_path)
    configure_metric_safety_fixture(product)
    registry = tmp_path / "registry"

    with pytest.raises(ProviderExecutionError) as captured:
        process_product(
            product,
            registry_root=registry,
            provider=UnsafeMetricProvider(),
        )

    assert captured.value.code == "LLM_METRIC_SQL_UNSAFE"
    assert captured.value.kind is ProviderFailureKind.OUTPUT
    assert "DELETE" not in str(captured.value)
    assert not (product / "generated").exists()
    assert not registry.exists()


def test_pipeline_classifies_duplicate_metric_names_as_provider_output(
    tmp_path: Path,
) -> None:
    product = create_product_fixture(tmp_path)
    configure_metric_safety_fixture(product)
    registry = tmp_path / "registry"

    with pytest.raises(ProviderExecutionError) as captured:
        process_product(
            product,
            registry_root=registry,
            provider=DuplicateMetricProvider(),
        )

    assert captured.value.code == "LLM_METRIC_NAME_DUPLICATE"
    assert captured.value.kind is ProviderFailureKind.OUTPUT
    assert not (product / "generated").exists()
    assert not registry.exists()


def test_reprocessing_removes_newly_unsupported_metric_from_current_registry(
    tmp_path: Path,
) -> None:
    product = create_product_fixture(tmp_path)
    configure_metric_safety_fixture(product)
    registry = tmp_path / "registry"
    process_product(
        product,
        registry_root=registry,
        provider=InitiallyAcceptedMetricsProvider(),
    )
    config_path = product / "product.yaml"
    config = yaml.safe_load(config_path.read_text(encoding="utf-8"))
    config["operation"] = "update"
    config["base_version"] = 1
    config["version"] = 2
    config["description"] = "Updated campaign and sales performance product."
    config_path.write_text(
        yaml.safe_dump(config, sort_keys=False),
        encoding="utf-8",
    )

    process_product(
        product,
        registry_root=registry,
        provider=DatasetSafetyProvider(),
    )

    product_record = json.loads(
        (registry / "products" / f"{PRODUCT_ID}.json").read_text()
    )
    assert [metric["name"] for metric in product_record["metrics"]] == [
        "Campaign Count"
    ]


def test_pipeline_builds_stable_metric_and_fk_relationship_ids(tmp_path: Path) -> None:
    product = create_product_fixture(tmp_path)
    dictionary_path = product / "sources" / "dictionary" / "dictionary.xlsx"
    workbook = Workbook()
    sheet = workbook.active
    sheet.append(
        [
            "platform",
            "catalog",
            "schema",
            "table",
            "column",
            "data_type",
            "nullable",
            "pk",
            "fk_table",
            "fk_column",
        ]
    )
    sheet.append(
        ["erp", "analytics", "sales", "customers", "customer_id", "INT64", False, True]
    )
    sheet.append(["erp", "analytics", "sales", "orders", "order_id", "INT64", False, True])
    sheet.append(
        [
            "erp",
            "analytics",
            "sales",
            "orders",
            "customer_id",
            "INT64",
            False,
            False,
            "customers",
            "customer_id",
        ]
    )
    workbook.save(dictionary_path)
    config = yaml.safe_load((product / "product.yaml").read_text(encoding="utf-8"))
    config["tables"].append(
        {
            "locator": "erp|analytics|sales|customers",
            "table_id": CUSTOMERS_TABLE_ID,
            "version": 1,
            "usage": "REFERENCE",
        }
    )
    (product / "product.yaml").write_text(
        yaml.safe_dump(config, sort_keys=False), encoding="utf-8"
    )
    registry = tmp_path / "registry"

    process_product(product, registry_root=registry, provider=FakeMetricProvider())
    first = json.loads((product / "generated" / "ossie-model.json").read_text())
    first_record = json.loads((registry / "products" / f"{PRODUCT_ID}.json").read_text())

    process_product(product, registry_root=registry, provider=FakeMetricProvider())
    second = json.loads((product / "generated" / "ossie-model.json").read_text())
    second_record = json.loads((registry / "products" / f"{PRODUCT_ID}.json").read_text())

    model = first["semantic_model"][0]
    assert model["metrics"][0]["name"] == "order_count"
    assert model["relationships"][0]["from_columns"] == ["customer_id"]
    assert first == second
    assert first_record["metrics"][0]["metric_id"] == second_record["metrics"][0]["metric_id"]
    assert (
        first_record["relationships"][0]["relationship_id"]
        == second_record["relationships"][0]["relationship_id"]
    )
