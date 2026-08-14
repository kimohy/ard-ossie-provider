from __future__ import annotations

import ctypes
import hashlib
from pathlib import Path
from types import SimpleNamespace
from typing import Any

import pytest

from ard_ossie.canonical import canonical_hash
from ard_ossie.docling_parser import DoclingParser, Evidence
from ard_ossie.ingestion import SourceFile, SourceRole
from ard_ossie.semantic import parser as semantic_parser
from ard_ossie.semantic.correction import OcrCorrectionApplication
from ard_ossie.semantic.models import (
    ExtractionMode,
    NativeDocument,
    NativeGroup,
    NativeTable,
    NativeTableCell,
    ParagraphBlock,
    RepairBlock,
    RepairCell,
    RepairPlan,
    SemanticStructureRepairRecord,
    SourceBox,
    SourceSpan,
    TableBlock,
    TableCellBlock,
    make_span_id,
)
from ard_ossie.semantic.parser import _new_full_page_ocr_converter
from ard_ossie.semantic.repair import (
    REPAIR_PROMPT_VERSION,
    RepairApplication,
    SemanticStructureRepairPlanner,
    semantic_structure_repair_schema,
)
from ard_ossie.semantic.sources import extract_docx_native
from ard_ossie.semantic.structure import (
    StructureBlock,
    StructureCell,
    StructureDocument,
    StructureTable,
    reconcile_structure,
)
from tests.unit.semantic.test_docx_source import docx_source, xml_document
from tests.unit.semantic.test_pdf_source import (
    FakePdfDocument,
    FakePdfium,
    FakePdfiumError,
    FakePdfPage,
    FakeTextPage,
    semantic_pdf_source,
)

CONTROLLED_BYTES = b"controlled"
CONTROLLED_HASH = hashlib.sha256(CONTROLLED_BYTES).hexdigest()


class FakeDocument:
    def export_to_markdown(self) -> str:
        return "# Sales Order\n\nNet revenue excludes tax."

    def iterate_items(self):
        bbox = SimpleNamespace(l=10.0, t=20.0, r=100.0, b=40.0)
        provenance = SimpleNamespace(page_no=2, bbox=bbox, charspan=(0, 31))
        item = SimpleNamespace(text="Net revenue excludes tax.", prov=[provenance])
        yield item, 1


class FakeHtmlDocument:
    def export_to_markdown(self) -> str:
        return "# 제품 개요\n\n사용자가 입력한 제품 목적"

    def iterate_items(self):
        item = SimpleNamespace(text="사용자가 입력한 제품 목적", prov=[])
        yield item, 2


class FakeAiGeneratedHtmlDocument:
    def export_to_markdown(self) -> str:
        return "# 제품 개요\n\n(AI 자동생성) 데이터 요약\n\n자동 요약 값\n\n사용자 설명"

    def iterate_items(self):
        for text in (
            "(AI 자동생성) 데이터 요약",
            "자동 요약 값",
            "사용자 설명",
        ):
            yield SimpleNamespace(text=text, prov=[]), 5


class FakeConverter:
    def __init__(self, document: object | None = None) -> None:
        self.converted_paths: list[str] = []
        self.document = document or FakeDocument()

    def convert(self, source: str) -> object:
        self.converted_paths.append(source)
        return SimpleNamespace(document=self.document)


class FailingConversionConverter:
    def __init__(self) -> None:
        self.converted_paths: list[str] = []

    def convert(self, source: str) -> object:
        from docling.exceptions import ConversionError

        self.converted_paths.append(source)
        raise ConversionError("fixture conversion failure")


class FakeBox:
    def __init__(self, left: float, bottom: float, right: float, top: float) -> None:
        self.l = left * 100
        self.b = (1 - top) * 200
        self.r = right * 100
        self.t = (1 - bottom) * 200

    def to_bottom_left_origin(self, page_height: float) -> SimpleNamespace:
        return SimpleNamespace(
            l=self.l,
            b=page_height - self.t,
            r=self.r,
            t=page_height - self.b,
        )


class SegmentedTextPage(FakeTextPage):
    def __init__(self, segments: list[tuple[str, tuple[float, float, float, float]]]) -> None:
        super().__init__("".join(text for text, _box in segments))
        self._segments = segments
        self._offsets: list[tuple[int, int, tuple[float, float, float, float]]] = []
        self._raw_handles: dict[int, ctypes._Pointer[ctypes.c_int]] = {}
        cursor = 0
        for text, box in segments:
            self._offsets.append((cursor, cursor + len(text), box))
            cursor += len(text)

    def get_textobj(self, index: int) -> object:
        for segment_index, (start, end, _box) in enumerate(self._offsets):
            if start <= index < end:
                raw = self._raw_handles.setdefault(
                    segment_index,
                    ctypes.pointer(ctypes.c_int(segment_index + 1)),
                )
                return SimpleNamespace(raw=raw)
        raise AssertionError(index)

    def get_charbox(self, index: int) -> tuple[float, float, float, float]:
        for start, end, box in self._offsets:
            if start <= index < end:
                left, bottom, right, top = box
                width = (right - left) / (end - start)
                offset = index - start
                return (
                    (left + width * offset) * 100,
                    bottom * 200,
                    (left + width * (offset + 1)) * 100,
                    top * 200,
                )
        raise AssertionError(index)


class SegmentedPdfPage(FakePdfPage):
    def __init__(self, segments: list[tuple[str, tuple[float, float, float, float]]]) -> None:
        super().__init__("")
        self._segments = segments

    def get_textpage(self) -> SegmentedTextPage:
        self.text_page = SegmentedTextPage(self._segments)
        return self.text_page


class TitleItem:
    def __init__(self, text: str, box: tuple[float, float, float, float]) -> None:
        self.orig = text
        self.text = f"Docling hint: {text}"
        self.prov = [SimpleNamespace(page_no=1, bbox=FakeBox(*box))]


class ListItem(TitleItem):
    enumerated = False


class TextItem(TitleItem):
    pass


class SectionHeaderItem(TitleItem):
    level = 1


class TableItem(TitleItem):
    def __init__(self, cells: list[object]) -> None:
        super().__init__("table hint", (0.05, 0.15, 0.90, 0.40))
        self.data = SimpleNamespace(num_rows=2, num_cols=2, table_cells=cells)


def _cell(
    text: str,
    row: int,
    column: int,
    box: tuple[float, float, float, float],
) -> object:
    return SimpleNamespace(
        text=text,
        start_row_offset_idx=row,
        end_row_offset_idx=row + 1,
        start_col_offset_idx=column,
        end_col_offset_idx=column + 1,
        column_header=row == 0,
        bbox=FakeBox(*box),
    )


def structured_pdf_text() -> list[tuple[str, tuple[float, float, float, float]]]:
    return [
        ("개인정보", (0.05, 0.80, 0.20, 0.90)),
        ("목록항목", (0.05, 0.65, 0.20, 0.72)),
        ("설명문단", (0.05, 0.50, 0.30, 0.58)),
        ("항목", (0.05, 0.30, 0.45, 0.40)),
        ("값", (0.50, 0.30, 0.90, 0.40)),
        ("유형", (0.05, 0.15, 0.45, 0.25)),
        ("필수", (0.50, 0.15, 0.90, 0.25)),
    ]


def structured_pdf_document() -> object:
    cells = [
        _cell("항목", 0, 0, (0.05, 0.30, 0.45, 0.40)),
        _cell("값", 0, 1, (0.50, 0.30, 0.90, 0.40)),
        _cell("유형", 1, 0, (0.05, 0.15, 0.45, 0.25)),
        _cell("필수", 1, 1, (0.50, 0.15, 0.90, 0.25)),
    ]
    items = [
        (TitleItem("개 인정보", (0.05, 0.80, 0.20, 0.90)), 1),
        (ListItem("목록 항목", (0.05, 0.65, 0.20, 0.72)), 1),
        (TextItem("설명 문단", (0.05, 0.50, 0.30, 0.58)), 1),
        (TableItem(cells), 1),
    ]
    return SimpleNamespace(
        pages={1: SimpleNamespace(size=SimpleNamespace(width=100.0, height=200.0))},
        iterate_items=lambda: iter(items),
    )


def structured_ocr_document() -> object:
    items = [
        (TextItem("OCR_PAGE_ONE", (0.05, 0.70, 0.90, 0.80)), 1),
        (TextItem("OCR_PAGE_TWO", (0.05, 0.50, 0.90, 0.60)), 1),
    ]
    items[1][0].prov = [SimpleNamespace(page_no=2, bbox=FakeBox(0.05, 0.50, 0.90, 0.60))]
    return SimpleNamespace(
        pages={
            1: SimpleNamespace(size=SimpleNamespace(width=100.0, height=200.0)),
            2: SimpleNamespace(size=SimpleNamespace(width=100.0, height=200.0)),
        },
        iterate_items=lambda: iter(items),
    )


class FailingIfCalledPlanner:
    def __init__(self) -> None:
        self.call_count = 0

    def repair(self, *_args: object, **_kwargs: object) -> RepairApplication:
        self.call_count += 1
        raise AssertionError("repair planner was unexpectedly called")


class ApplyingOcrCorrectionPlanner:
    def __init__(self) -> None:
        self.call_count = 0

    def correct(self, _source: SourceFile, native: NativeDocument, **_kwargs: object):
        self.call_count += 1
        first = native.spans[0]
        corrected_text = "개인정보"
        corrected = first.model_copy(
            update={
                "text": corrected_text,
                "text_hash": hashlib.sha256(corrected_text.encode("utf-8")).hexdigest(),
            }
        )
        return OcrCorrectionApplication(
            document=native.model_copy(update={"spans": (corrected, *native.spans[1:])}),
            audits=(),
            warning_codes=(),
        )


class RejectingPlanner:
    def __init__(self) -> None:
        self.call_count = 0

    def repair(self, native: Any, *_args: object, **_kwargs: object) -> RepairApplication:
        self.call_count += 1
        plan = RepairPlan(blocks=[])
        return RepairApplication(
            blocks=(),
            record=SemanticStructureRepairRecord(
                source_hash=native.source_hash,
                ordered_span_hashes=[span.text_hash for span in native.spans],
                parser_version="semantic-structure-v1",
                prompt_version="semantic-structure-repair-v1",
                schema_hash="b" * 64,
                provider="fake",
                model="fake",
                outcome="rejected",
                plan=plan,
                provider_error_code=None,
                validation_codes=["SEMANTIC_REPAIR_SCHEMA_INVALID"],
                applied_orders=[],
                rejected_orders=[],
                plan_hash=canonical_hash(plan.model_dump(mode="json")),
            ),
        )


class DetailedRejectingPlanner:
    def repair(self, native: Any, *_args: object, **_kwargs: object) -> RepairApplication:
        plan = RepairPlan(blocks=[])
        return RepairApplication(
            blocks=(),
            record=SemanticStructureRepairRecord(
                source_hash=native.source_hash,
                ordered_span_hashes=[span.text_hash for span in native.spans],
                parser_version="semantic-structure-v1",
                prompt_version=REPAIR_PROMPT_VERSION,
                schema_hash=canonical_hash(semantic_structure_repair_schema()),
                provider="fake",
                model="fake",
                outcome="rejected",
                plan=plan,
                provider_error_code=None,
                validation_codes=[
                    "SEMANTIC_REPAIR_MISSING_SPAN",
                    "SEMANTIC_REPAIR_ORDER_INVALID",
                ],
                applied_orders=[],
                rejected_orders=[],
                plan_hash=canonical_hash(plan.model_dump(mode="json")),
            ),
        )


class ApplyingTablePlanner:
    def repair(self, native: Any, *_args: object, **_kwargs: object) -> RepairApplication:
        table = native.tables[0]
        plan = RepairPlan(
            blocks=[
                RepairBlock(
                    kind="table",
                    order=10,
                    span_ids=[],
                    heading_level=None,
                    list_kind=None,
                    list_depth=None,
                    row_count=table.row_count,
                    column_count=table.column_count,
                    cells=[
                        RepairCell(
                            start_row=cell.start_row,
                            end_row=cell.end_row,
                            start_column=cell.start_column,
                            end_column=cell.end_column,
                            span_ids=list(cell.span_ids),
                            column_header=cell.column_header,
                        )
                        for cell in table.cells
                    ],
                    exclusion_kind=None,
                    confidence=1.0,
                )
            ]
        )
        block = TableBlock(
            order=10,
            row_count=table.row_count,
            column_count=table.column_count,
            cells=tuple(
                TableCellBlock(
                    start_row=cell.start_row,
                    end_row=cell.end_row,
                    start_column=cell.start_column,
                    end_column=cell.end_column,
                    span_ids=cell.span_ids,
                    column_header=cell.column_header,
                )
                for cell in table.cells
            ),
        )
        return RepairApplication(
            blocks=(block,),
            record=SemanticStructureRepairRecord(
                source_hash=native.source_hash,
                ordered_span_hashes=[span.text_hash for span in native.spans],
                parser_version="semantic-structure-v1",
                prompt_version="semantic-structure-repair-v1",
                schema_hash="b" * 64,
                provider="fake",
                model="fake",
                outcome="applied",
                plan=plan,
                provider_error_code=None,
                validation_codes=[],
                applied_orders=[10],
                rejected_orders=[],
                plan_hash=canonical_hash(plan.model_dump(mode="json")),
            ),
        )


class ProviderFailingPlanner:
    def repair(self, *_args: object, **_kwargs: object) -> RepairApplication:
        from ard_ossie.llm import ProviderExecutionError, ProviderFailureKind

        raise ProviderExecutionError(
            "LLM_PROVIDER_TRANSIENT_FAILED",
            kind=ProviderFailureKind.TRANSIENT,
        )


class FixedRepairPlanner:
    def __init__(self, plan: RepairPlan, blocks: tuple[object, ...]) -> None:
        self.plan = plan
        self.blocks = blocks

    def repair(self, native: Any, *_args: object, **_kwargs: object) -> RepairApplication:
        return RepairApplication(
            blocks=self.blocks,  # type: ignore[arg-type]
            record=SemanticStructureRepairRecord(
                source_hash=native.source_hash,
                ordered_span_hashes=[span.text_hash for span in native.spans],
                parser_version="semantic-structure-v1",
                prompt_version=REPAIR_PROMPT_VERSION,
                schema_hash=canonical_hash(semantic_structure_repair_schema()),
                provider="fake",
                model="fake",
                outcome="applied",
                plan=self.plan,
                provider_error_code=None,
                validation_codes=[],
                applied_orders=[block.order for block in self.plan.blocks],
                rejected_orders=[],
                plan_hash=canonical_hash(self.plan.model_dump(mode="json")),
            ),
        )


class NoCallProvider:
    def capabilities(self) -> dict[str, str]:
        return {"provider": "fake", "model": "fake"}

    def generate_structured(self, **_kwargs: object) -> dict[str, object]:
        raise AssertionError("trusted repair should be reused")


def controlled_source(tmp_path: Path) -> SourceFile:
    path = tmp_path / "controlled.docx"
    contents = CONTROLLED_BYTES
    path.write_bytes(contents)
    return SourceFile(
        role=SourceRole.SEMANTIC_DOCUMENT,
        path=path,
        relative_path="semantic/controlled.docx",
        sha256=CONTROLLED_HASH,
        size_bytes=10,
        snapshot=contents,
    )


def controlled_span(ordinal: int, text: str) -> SourceSpan:
    return SourceSpan(
        span_id=make_span_id(CONTROLLED_HASH, ordinal),
        ordinal=ordinal,
        text=text,
        text_hash=hashlib.sha256(text.encode()).hexdigest(),
    )


def controlled_native_table(
    texts: tuple[str, ...],
    *,
    row_count: int,
    column_count: int,
    table_ordinals: tuple[int, ...] | None = None,
) -> NativeDocument:
    spans = tuple(controlled_span(index, text) for index, text in enumerate(texts))
    selected = table_ordinals if table_ordinals is not None else tuple(range(len(spans)))
    table_ids = tuple(spans[index].span_id for index in selected)
    cells = tuple(
        NativeTableCell(
            start_row=index // column_count,
            end_row=index // column_count + 1,
            start_column=index % column_count,
            end_column=index % column_count + 1,
            span_ids=((span_id,) if span_id else ()),
            column_header=index // column_count == 0,
        )
        for index, span_id in enumerate(table_ids)
    )
    groups = [NativeGroup(order=0, kind="table", span_ids=table_ids, table_index=0)]
    groups.extend(
        NativeGroup(order=order, kind="paragraph", span_ids=(span.span_id,))
        for order, span in enumerate(spans, start=1)
        if span.ordinal not in selected
    )
    return NativeDocument(
        source_hash=CONTROLLED_HASH,
        extraction_mode=ExtractionMode.DOCX_XML,
        page_count=0,
        parser_versions={},
        spans=spans,
        groups=tuple(groups),
        tables=(
            NativeTable(
                order=0,
                row_count=row_count,
                column_count=column_count,
                cells=cells,
            ),
        ),
    )


def controlled_native_paragraphs(texts: tuple[str, ...]) -> NativeDocument:
    spans = tuple(controlled_span(index, text) for index, text in enumerate(texts))
    return NativeDocument(
        source_hash=CONTROLLED_HASH,
        extraction_mode=ExtractionMode.DOCX_XML,
        page_count=0,
        parser_versions={},
        spans=spans,
        groups=tuple(
            NativeGroup(order=index, kind="paragraph", span_ids=(span.span_id,))
            for index, span in enumerate(spans)
        ),
        tables=(),
    )


def controlled_ocr_native(texts: tuple[str, ...]) -> NativeDocument:
    spans = tuple(
        controlled_span(index, text).model_copy(
            update={
                "page": 1,
                "bbox": SourceBox(
                    left=0.05,
                    bottom=0.80 - index * 0.10,
                    right=0.95,
                    top=0.88 - index * 0.10,
                ),
            }
        )
        for index, text in enumerate(texts)
    )
    return NativeDocument(
        source_hash=CONTROLLED_HASH,
        extraction_mode=ExtractionMode.OCR,
        page_count=1,
        parser_versions={"ocr": "fixture-v1"},
        spans=spans,
        groups=(),
        tables=(),
    )


def parse_controlled(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
    native: NativeDocument,
    *,
    planner: object | None = None,
    trusted_record: SemanticStructureRepairRecord | None = None,
) -> object:
    monkeypatch.setattr(
        semantic_parser,
        "_native_and_structure",
        lambda *_args, **_kwargs: (native, StructureDocument(blocks=())),
    )
    return DoclingParser(
        structure_repair_planner=planner,  # type: ignore[arg-type]
        trusted_repair_record=trusted_record,
    ).parse(controlled_source(tmp_path))


def table_repair_block(
    native: NativeDocument,
    *,
    order: int = 10,
    row_count: int | None = None,
    column_count: int | None = None,
) -> tuple[RepairBlock, TableBlock]:
    table = native.tables[0]
    target_rows = row_count or table.row_count
    target_columns = column_count or table.column_count
    span_ids = tuple(span_id for cell in table.cells for span_id in cell.span_ids)
    repair_cells = [
        RepairCell(
            start_row=index // target_columns,
            end_row=index // target_columns + 1,
            start_column=index % target_columns,
            end_column=index % target_columns + 1,
            span_ids=[span_id],
            column_header=index // target_columns == 0,
        )
        for index, span_id in enumerate(span_ids)
    ]
    repair = RepairBlock(
        kind="table",
        order=order,
        span_ids=[],
        heading_level=None,
        list_kind=None,
        list_depth=None,
        row_count=target_rows,
        column_count=target_columns,
        cells=repair_cells,
        exclusion_kind=None,
        confidence=1.0,
    )
    semantic = TableBlock(
        order=order,
        row_count=target_rows,
        column_count=target_columns,
        cells=tuple(
            TableCellBlock(
                start_row=cell.start_row,
                end_row=cell.end_row,
                start_column=cell.start_column,
                end_column=cell.end_column,
                span_ids=tuple(cell.span_ids),
                column_header=cell.column_header,
            )
            for cell in repair_cells
        ),
    )
    return repair, semantic


def test_semantic_pdf_prefers_embedded_text_and_still_applies_visual_correction(
    tmp_path: Path,
) -> None:
    source = semantic_pdf_source(tmp_path)
    converter = FakeConverter(structured_pdf_document())
    full_page_ocr_converter = FakeConverter(structured_pdf_document())
    pdfium = FakePdfium(
        document=FakePdfDocument(pages=[SegmentedPdfPage(structured_pdf_text())])
    )
    planner = FailingIfCalledPlanner()
    correction_planner = ApplyingOcrCorrectionPlanner()

    parsed = DoclingParser(
        converter=converter,
        full_page_ocr_converter=full_page_ocr_converter,
        pdfium=pdfium,
        structure_repair_planner=planner,
        ocr_correction_planner=correction_planner,
    ).parse(source)

    assert parsed.markdown == (
        "# 개인정보\n\n"
        "- 목록항목\n\n"
        "설명문단\n\n"
        "| 항목 | 값 |\n"
        "| --- | --- |\n"
        "| 유형 | 필수 |\n"
    )
    assert "개 인정보" not in parsed.markdown
    assert parsed.evidence[0].locator == {
        "document": "semantic/semantic.pdf",
        "span_id": parsed.evidence[0].locator["span_id"],
        "page": 1,
        "order": 0,
        "bbox": {"left": 0.05, "bottom": 0.8, "right": 0.2, "top": 0.9},
    }
    assert [item.excerpt for item in parsed.evidence] == [
        "개인정보",
        "목록항목",
        "설명문단",
        "항목",
        "값",
        "유형",
        "필수",
    ]
    assert [Path(item).suffix for item in converter.converted_paths] == [
        source.path.suffix
    ]
    assert all(
        Path(item) != source.path and not Path(item).exists()
        for item in converter.converted_paths
    )
    assert full_page_ocr_converter.converted_paths == []
    assert parsed.semantic_fidelity is not None
    assert parsed.semantic_fidelity.source_text_coverage == 1.0
    assert parsed.semantic_fidelity.extraction_mode is ExtractionMode.PDF_EMBEDDED
    assert parsed.semantic_fidelity.status == "PASS"
    assert parsed.semantic_fidelity.paragraph_count == 1
    assert parsed.semantic_fidelity.degraded_block_count == 0
    assert parsed.semantic_fidelity.table_results[0].matched_cell_count == 4
    assert parsed.semantic_fidelity.table_results[0].total_cell_count == 4
    assert planner.call_count == 0
    assert correction_planner.call_count == 1


@pytest.mark.parametrize(
    "pdfium",
    [
        pytest.param(
            FakePdfium(["EMBEDDED_PAGE_ONE", "   "]),
            id="empty-normalized-page",
        ),
        pytest.param(
            FakePdfium(
                document=FakePdfDocument(
                    pages=[
                        FakePdfPage(
                            "EMBEDDED_PAGE_ONE",
                            get_textpage_error=FakePdfiumError("text page unavailable"),
                        )
                    ]
                )
            ),
            id="get-textpage-error",
        ),
        pytest.param(
            FakePdfium(
                document=FakePdfDocument(
                    pages=[
                        FakePdfPage(
                            "EMBEDDED_PAGE_ONE",
                            get_text_range_error=FakePdfiumError("text unavailable"),
                        )
                    ]
                )
            ),
            id="get-text-range-error",
        ),
    ],
)
def test_partial_pdf_uses_one_whole_document_full_page_ocr_catalog(
    tmp_path: Path,
    pdfium: FakePdfium,
) -> None:
    converter = FakeConverter(structured_pdf_document())
    full_page_ocr_converter = FakeConverter(structured_ocr_document())

    parsed = DoclingParser(
        converter=converter,
        full_page_ocr_converter=full_page_ocr_converter,
        pdfium=pdfium,
    ).parse(semantic_pdf_source(tmp_path))

    assert "OCR\\_PAGE\\_ONE" in parsed.markdown
    assert "OCR\\_PAGE\\_TWO" in parsed.markdown
    assert [item.excerpt for item in parsed.evidence] == ["OCR_PAGE_ONE", "OCR_PAGE_TWO"]
    assert "EMBEDDED_PAGE_ONE" not in parsed.markdown
    assert converter.converted_paths == []
    assert [Path(item).suffix for item in full_page_ocr_converter.converted_paths] == [".pdf"]
    assert all(not Path(item).exists() for item in full_page_ocr_converter.converted_paths)
    assert parsed.semantic_fidelity is not None
    assert parsed.semantic_fidelity.extraction_mode is ExtractionMode.OCR
    assert parsed.semantic_fidelity.status == "WARN"
    assert pdfium.document.closed is True


def test_docling_adapter_uses_full_page_ocr_when_embedded_pdf_cannot_open(
    tmp_path: Path,
) -> None:
    converter = FakeConverter(structured_pdf_document())
    full_page_ocr_converter = FakeConverter(structured_ocr_document())
    pdfium = FakePdfium(document_open_error=True)

    parsed = DoclingParser(
        converter=converter,
        full_page_ocr_converter=full_page_ocr_converter,
        pdfium=pdfium,
    ).parse(semantic_pdf_source(tmp_path))

    assert "OCR\\_PAGE\\_ONE" in parsed.markdown
    assert converter.converted_paths == []
    assert len(full_page_ocr_converter.converted_paths) == 1
    assert pdfium.document is None


def test_failed_table_repair_emits_lossless_block_and_degraded_fidelity(
    tmp_path: Path,
) -> None:
    from docx import Document

    path = tmp_path / "semantic.docx"
    document = Document()
    table = document.add_table(rows=2, cols=2)
    for cell, text in zip(
        (cell for row in table.rows for cell in row.cells),
        ("항목", "값", "유형", "필수"),
        strict=True,
    ):
        cell.text = text
    document.save(path)
    source = SourceFile(
        role=SourceRole.SEMANTIC_DOCUMENT,
        path=path,
        relative_path="semantic/semantic.docx",
        sha256=hashlib.sha256(path.read_bytes()).hexdigest(),
        size_bytes=path.stat().st_size,
        snapshot=path.read_bytes(),
    )
    planner = RejectingPlanner()

    parsed = DoclingParser(
        converter=FakeConverter(SimpleNamespace(iterate_items=lambda: iter(()))),
        structure_repair_planner=planner,
    ).parse(source)

    assert not parsed.markdown.startswith("<pre>")
    assert "항목값유형필수" in parsed.markdown
    assert parsed.semantic_fidelity is not None
    assert parsed.semantic_fidelity.status == "WARN"
    assert parsed.semantic_fidelity.degraded_block_count == 1
    assert parsed.semantic_fidelity.degraded_blocks[0].reason == "repair_rejected"
    assert parsed.semantic_fidelity.table_results[0].status == "degraded"
    assert parsed.semantic_repair is not None
    assert parsed.semantic_repair.outcome == "rejected"
    assert planner.call_count == 1
    dumped = parsed.model_dump(mode="json")
    assert "semantic_fidelity" not in dumped
    assert "semantic_repair" not in dumped


def test_repaired_table_status_survives_final_source_order_normalization(
    tmp_path: Path,
) -> None:
    from docx import Document

    path = tmp_path / "semantic.docx"
    document = Document()
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "값"
    document.save(path)
    source = SourceFile(
        role=SourceRole.SEMANTIC_DOCUMENT,
        path=path,
        relative_path="semantic/semantic.docx",
        sha256=hashlib.sha256(path.read_bytes()).hexdigest(),
        size_bytes=path.stat().st_size,
        snapshot=path.read_bytes(),
    )

    parsed = DoclingParser(
        converter=FakeConverter(SimpleNamespace(iterate_items=lambda: iter(()))),
        structure_repair_planner=ApplyingTablePlanner(),
    ).parse(source)

    assert parsed.semantic_fidelity is not None
    assert parsed.semantic_fidelity.table_results[0].order == 0
    assert parsed.semantic_fidelity.table_results[0].status == "repaired"


def test_native_table_spans_repaired_as_paragraph_degrade_atomically(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    native = controlled_native_table(("A", "C"), row_count=2, column_count=1)
    span_ids = [span.span_id for span in native.spans]
    repair = RepairBlock(
        kind="paragraph",
        order=10,
        span_ids=span_ids,
        heading_level=None,
        list_kind=None,
        list_depth=None,
        row_count=None,
        column_count=None,
        cells=[],
        exclusion_kind=None,
        confidence=1.0,
    )
    plan = RepairPlan(blocks=[repair])

    parsed = parse_controlled(
        tmp_path,
        monkeypatch,
        native,
        planner=FixedRepairPlanner(
            plan,
            (ParagraphBlock(order=10, span_ids=tuple(span_ids)),),
        ),
    )

    assert parsed.markdown == "AC\n"
    assert parsed.semantic_fidelity.status == "WARN"
    assert parsed.semantic_fidelity.table_count == 1
    assert parsed.semantic_fidelity.table_results[0].status == "degraded"
    assert parsed.semantic_repair.outcome == "rejected"
    assert parsed.semantic_repair.applied_orders == []
    assert parsed.semantic_repair.rejected_orders == [
        block.order for block in parsed.semantic_repair.plan.blocks
    ]
    assert "SEMANTIC_REPAIR_TABLE_INVALID" in parsed.semantic_repair.validation_codes
    assert parsed.semantic_repair.plan_hash == canonical_hash(
        parsed.semantic_repair.plan.model_dump(mode="json")
    )


def test_partial_native_table_repair_rejects_every_touching_block(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    native = controlled_native_table(("A", "C"), row_count=2, column_count=1)
    first, second = (span.span_id for span in native.spans)
    table_repair = RepairBlock(
        kind="table",
        order=10,
        span_ids=[],
        heading_level=None,
        list_kind=None,
        list_depth=None,
        row_count=1,
        column_count=1,
        cells=[
            RepairCell(
                start_row=0,
                end_row=1,
                start_column=0,
                end_column=1,
                span_ids=[first],
                column_header=True,
            )
        ],
        exclusion_kind=None,
        confidence=1.0,
    )
    paragraph_repair = RepairBlock(
        kind="paragraph",
        order=11,
        span_ids=[second],
        heading_level=None,
        list_kind=None,
        list_depth=None,
        row_count=None,
        column_count=None,
        cells=[],
        exclusion_kind=None,
        confidence=1.0,
    )
    plan = RepairPlan(blocks=[table_repair, paragraph_repair])
    blocks = (
        TableBlock(
            order=10,
            row_count=1,
            column_count=1,
            cells=(
                TableCellBlock(
                    start_row=0,
                    end_row=1,
                    start_column=0,
                    end_column=1,
                    span_ids=(first,),
                    column_header=True,
                ),
            ),
        ),
        ParagraphBlock(order=11, span_ids=(second,)),
    )

    parsed = parse_controlled(
        tmp_path,
        monkeypatch,
        native,
        planner=FixedRepairPlanner(plan, blocks),
    )

    assert parsed.markdown == "AC\n"
    assert parsed.semantic_fidelity.status == "WARN"
    assert parsed.semantic_fidelity.table_results[0].status == "degraded"
    assert parsed.semantic_repair.applied_orders == []
    assert parsed.semantic_repair.rejected_orders == [
        block.order for block in parsed.semantic_repair.plan.blocks
    ]


def test_wrong_grid_native_table_repair_degrades_complete_table(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    native = controlled_native_table(("A", "C"), row_count=2, column_count=1)
    repair, semantic = table_repair_block(
        native,
        row_count=1,
        column_count=2,
    )
    plan = RepairPlan(blocks=[repair])

    parsed = parse_controlled(
        tmp_path,
        monkeypatch,
        native,
        planner=FixedRepairPlanner(plan, (semantic,)),
    )

    assert parsed.markdown == "AC\n"
    assert parsed.semantic_fidelity.status == "WARN"
    assert parsed.semantic_fidelity.table_results[0].status == "degraded"
    assert parsed.semantic_repair.outcome == "rejected"
    assert "SEMANTIC_REPAIR_TABLE_INVALID" in parsed.semantic_repair.validation_codes


def test_valid_native_table_repair_remaps_record_and_remains_reusable(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    native = controlled_native_table(("A", "C"), row_count=2, column_count=1)
    repair, semantic = table_repair_block(native, order=10)
    plan = RepairPlan(blocks=[repair])

    parsed = parse_controlled(
        tmp_path,
        monkeypatch,
        native,
        planner=FixedRepairPlanner(plan, (semantic,)),
    )

    assert parsed.semantic_fidelity.status == "PASS"
    assert parsed.semantic_fidelity.table_results[0].status == "repaired"
    assert parsed.semantic_fidelity.table_results[0].order == 0
    assert parsed.semantic_repair.applied_orders == [0]
    assert parsed.semantic_repair.rejected_orders == []
    assert parsed.semantic_repair.plan.blocks[0].order == 0
    assert parsed.semantic_repair.plan_hash == canonical_hash(
        parsed.semantic_repair.plan.model_dump(mode="json")
    )

    reused = parse_controlled(
        tmp_path,
        monkeypatch,
        native,
        planner=SemanticStructureRepairPlanner(NoCallProvider()),  # type: ignore[arg-type]
        trusted_record=parsed.semantic_repair,
    )

    assert reused.semantic_repair.outcome == "reused"
    assert reused.semantic_repair.plan_hash == parsed.semantic_repair.plan_hash
    assert reused.semantic_fidelity.table_results[0].status == "repaired"


def test_blank_native_table_is_rendered_and_audited_once(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    native = NativeDocument(
        source_hash="a" * 64,
        extraction_mode=ExtractionMode.DOCX_XML,
        page_count=0,
        parser_versions={},
        spans=(),
        groups=(NativeGroup(order=0, kind="table", span_ids=(), table_index=0),),
        tables=(
            NativeTable(
                order=0,
                row_count=1,
                column_count=2,
                cells=(
                    NativeTableCell(
                        start_row=0,
                        end_row=1,
                        start_column=0,
                        end_column=1,
                        column_header=True,
                    ),
                    NativeTableCell(
                        start_row=0,
                        end_row=1,
                        start_column=1,
                        end_column=2,
                        column_header=True,
                    ),
                ),
            ),
        ),
    )

    parsed = parse_controlled(tmp_path, monkeypatch, native)

    assert parsed.markdown == "|  |  |\n| --- | --- |\n"
    assert parsed.semantic_fidelity.status == "PASS"
    assert parsed.semantic_fidelity.table_count == 1
    assert parsed.semantic_fidelity.row_count == 1
    assert parsed.semantic_fidelity.cell_count == 2
    assert len(parsed.semantic_fidelity.table_results) == 1


def test_crossing_native_table_region_degrades_in_exact_source_order(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    native = controlled_native_table(
        ("A", "AUX", "C"),
        row_count=1,
        column_count=2,
        table_ordinals=(0, 2),
    )

    parsed = parse_controlled(tmp_path, monkeypatch, native)

    assert parsed.markdown == "AAUXC\n"
    assert [item.excerpt for item in parsed.evidence] == ["A", "AUX", "C"]
    assert parsed.semantic_fidelity.status == "WARN"
    assert parsed.semantic_fidelity.degraded_block_count == 1
    assert parsed.semantic_fidelity.table_count == 1
    assert parsed.semantic_fidelity.table_results[0].status == "degraded"
    assert parsed.semantic_fidelity.preserved_span_count == 3


def test_nested_docx_deterministic_table_matches_preserve_every_span_once(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    source = docx_source(
        tmp_path,
        document=xml_document(
            """
<w:p><w:r><w:t>A</w:t></w:r></w:p>
<w:tbl><w:tblGrid><w:gridCol/></w:tblGrid><w:tr><w:tc>
 <w:p><w:r><w:t>before</w:t></w:r></w:p>
 <w:tbl><w:tblGrid><w:gridCol/></w:tblGrid><w:tr><w:tc>
  <w:p><w:r><w:t>nested</w:t></w:r></w:p>
 </w:tc></w:tr></w:tbl>
 <w:p><w:r><w:t>after</w:t></w:r></w:p>
</w:tc></w:tr></w:tbl>
<w:p><w:r><w:t>B</w:t></w:r></w:p>"""
        ),
    )
    native = extract_docx_native(source)
    outer = StructureBlock(
        kind="table",
        order=1,
        page=None,
        bbox=None,
        text_hint="beforeafter",
        table=StructureTable(
            row_count=1,
            column_count=1,
            cells=(StructureCell(0, 1, 0, 1, "beforeafter", False, None),),
        ),
    )
    nested = StructureBlock(
        kind="table",
        order=2,
        page=None,
        bbox=None,
        text_hint="nested",
        table=StructureTable(
            row_count=1,
            column_count=1,
            cells=(StructureCell(0, 1, 0, 1, "nested", False, None),),
        ),
    )
    skeleton = StructureDocument(blocks=(outer, nested))
    reconciled = reconcile_structure(native, skeleton)
    assert len(reconciled.blocks) == 2
    assert all(isinstance(block, TableBlock) for block in reconciled.blocks)
    monkeypatch.setattr(
        semantic_parser,
        "_native_and_structure",
        lambda *_args, **_kwargs: (native, skeleton),
    )

    parsed = DoclingParser().parse(source)

    assert [item.excerpt for item in parsed.evidence] == [
        "A",
        "before",
        "nested",
        "after",
        "B",
    ]
    assert parsed.markdown == "A\n\nbeforenestedafter\n\nB\n"
    assert parsed.semantic_fidelity.source_text_coverage == 1.0
    assert parsed.semantic_fidelity.unmatched_span_count == 0
    assert parsed.semantic_fidelity.duplicated_span_count == 0
    assert parsed.semantic_fidelity.table_count == 2
    assert [result.status for result in parsed.semantic_fidelity.table_results] == [
        "degraded",
        "degraded",
    ]


def test_crossing_table_deactivates_repairs_for_interleaved_auxiliary_span(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    native = controlled_native_table(
        ("A", "AUX", "C"),
        row_count=1,
        column_count=2,
        table_ordinals=(0, 2),
    )
    repair, semantic = table_repair_block(native, order=10)
    auxiliary_id = native.spans[1].span_id
    auxiliary_repair = RepairBlock(
        kind="paragraph",
        order=11,
        span_ids=[auxiliary_id],
        heading_level=None,
        list_kind=None,
        list_depth=None,
        row_count=None,
        column_count=None,
        cells=[],
        exclusion_kind=None,
        confidence=1.0,
    )
    plan = RepairPlan(blocks=[repair, auxiliary_repair])

    parsed = parse_controlled(
        tmp_path,
        monkeypatch,
        native,
        planner=FixedRepairPlanner(
            plan,
            (
                semantic,
                ParagraphBlock(order=11, span_ids=(auxiliary_id,)),
            ),
        ),
    )

    assert parsed.markdown == "AAUXC\n"
    assert parsed.semantic_repair.outcome == "rejected"
    assert parsed.semantic_repair.applied_orders == []
    assert parsed.semantic_repair.rejected_orders == [
        block.order for block in parsed.semantic_repair.plan.blocks
    ]


def test_ordinary_conversion_error_degrades_readable_docx_without_masking_text(
    tmp_path: Path,
) -> None:
    from docx import Document

    path = tmp_path / "semantic.docx"
    document = Document()
    document.add_paragraph("원문 유지")
    document.save(path)
    source = SourceFile(
        role=SourceRole.SEMANTIC_DOCUMENT,
        path=path,
        relative_path="semantic/semantic.docx",
        sha256=hashlib.sha256(path.read_bytes()).hexdigest(),
        size_bytes=path.stat().st_size,
        snapshot=path.read_bytes(),
    )

    parsed = DoclingParser(converter=FailingConversionConverter()).parse(source)

    assert parsed.markdown == "원문 유지\n"
    assert parsed.semantic_fidelity is not None
    assert parsed.semantic_fidelity.status == "WARN"
    assert parsed.semantic_fidelity.degraded_block_count == 1
    assert parsed.semantic_fidelity.degraded_blocks[0].reason == "provider_unavailable"
    assert parsed.semantic_fidelity.source_text_coverage == 1.0
    assert parsed.semantic_fidelity.unmatched_span_count == 0
    assert parsed.semantic_fidelity.duplicated_span_count == 0


def test_full_page_ocr_conversion_error_is_hard_unreadable(tmp_path: Path) -> None:
    from ard_ossie.semantic.sources import SemanticSourceError

    converter = FailingConversionConverter()

    with pytest.raises(SemanticSourceError, match="SEMANTIC_OCR_UNREADABLE") as exc_info:
        DoclingParser(
            converter=FakeConverter(structured_pdf_document()),
            full_page_ocr_converter=converter,
            pdfium=FakePdfium(["EMBEDDED_PAGE_ONE", ""]),
        ).parse(semantic_pdf_source(tmp_path))

    assert exc_info.value.code == "SEMANTIC_OCR_UNREADABLE"
    assert [Path(item).suffix for item in converter.converted_paths] == [".pdf"]
    assert all(not Path(item).exists() for item in converter.converted_paths)


def test_full_page_ocr_empty_native_catalog_is_hard_unreadable(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    import ard_ossie.semantic.parser as semantic_parser
    from ard_ossie.semantic.sources import SemanticSourceError

    source = semantic_pdf_source(tmp_path)
    monkeypatch.setattr(
        semantic_parser,
        "extract_ocr_native",
        lambda _source, _document: NativeDocument(
            source_hash=source.sha256,
            extraction_mode=ExtractionMode.OCR,
            page_count=1,
            parser_versions={},
            spans=(),
            groups=(),
            tables=(),
        ),
    )

    with pytest.raises(SemanticSourceError, match="SEMANTIC_OCR_UNREADABLE") as exc_info:
        DoclingParser(
            full_page_ocr_converter=FakeConverter(
                SimpleNamespace(iterate_items=lambda: iter(()), pages={})
            ),
            pdfium=FakePdfium(["EMBEDDED_PAGE_ONE", ""]),
        ).parse(source)

    assert exc_info.value.code == "SEMANTIC_OCR_UNREADABLE"


def test_provider_execution_failure_degrades_with_exact_native_table_text(tmp_path: Path) -> None:
    from docx import Document

    path = tmp_path / "semantic.docx"
    document = Document()
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "원문 유지"
    document.save(path)
    source = SourceFile(
        role=SourceRole.SEMANTIC_DOCUMENT,
        path=path,
        relative_path="semantic/semantic.docx",
        sha256=hashlib.sha256(path.read_bytes()).hexdigest(),
        size_bytes=path.stat().st_size,
        snapshot=path.read_bytes(),
    )

    parsed = DoclingParser(
        converter=FakeConverter(SimpleNamespace(iterate_items=lambda: iter(()))),
        structure_repair_planner=ProviderFailingPlanner(),
    ).parse(source)

    assert parsed.markdown == "원문 유지\n"
    assert parsed.semantic_fidelity is not None
    assert parsed.semantic_fidelity.status == "WARN"
    assert parsed.semantic_fidelity.degraded_blocks[0].reason == "provider_unavailable"


def test_provider_execution_failure_propagates_for_fail_fast_planner(tmp_path: Path) -> None:
    from docx import Document

    from ard_ossie.llm import ProviderExecutionError

    path = tmp_path / "semantic.docx"
    document = Document()
    document.add_paragraph("원문 유지")
    document.save(path)
    source = SourceFile(
        role=SourceRole.SEMANTIC_DOCUMENT,
        path=path,
        relative_path="semantic/semantic.docx",
        sha256=hashlib.sha256(path.read_bytes()).hexdigest(),
        size_bytes=path.stat().st_size,
        snapshot=path.read_bytes(),
    )
    planner = ProviderFailingPlanner()
    planner.propagate_provider_errors = True

    with pytest.raises(
        ProviderExecutionError,
        match="LLM_PROVIDER_TRANSIENT_FAILED",
    ):
        DoclingParser(
            converter=FakeConverter(SimpleNamespace(iterate_items=lambda: iter(()))),
            structure_repair_planner=planner,  # type: ignore[arg-type]
        ).parse(source)


def test_unresolved_docx_paragraphs_keep_native_paragraph_boundaries(
    tmp_path: Path,
) -> None:
    from docx import Document

    path = tmp_path / "semantic.docx"
    document = Document()
    document.add_paragraph("첫 문단")
    document.add_paragraph("둘째 문단")
    document.save(path)
    source = SourceFile(
        role=SourceRole.SEMANTIC_DOCUMENT,
        path=path,
        relative_path="semantic/semantic.docx",
        sha256=hashlib.sha256(path.read_bytes()).hexdigest(),
        size_bytes=path.stat().st_size,
        snapshot=path.read_bytes(),
    )

    parsed = DoclingParser(
        converter=FakeConverter(SimpleNamespace(iterate_items=lambda: iter(()))),
    ).parse(source)

    assert parsed.markdown == "첫 문단\n\n둘째 문단\n"
    assert parsed.semantic_fidelity is not None
    assert parsed.semantic_fidelity.status == "WARN"
    assert parsed.semantic_fidelity.paragraph_count == 0
    assert parsed.semantic_fidelity.degraded_block_count == 2
    assert [
        block.reason for block in parsed.semantic_fidelity.degraded_blocks
    ] == ["provider_unavailable", "provider_unavailable"]
    assert [item.excerpt for item in parsed.evidence] == ["첫 문단", "둘째 문단"]
    assert parsed.semantic_fidelity.preserved_span_count == 2
    assert parsed.semantic_fidelity.source_text_coverage == 1.0
    assert parsed.semantic_fidelity.unmatched_span_count == 0
    assert parsed.semantic_fidelity.duplicated_span_count == 0


@pytest.mark.parametrize(
    ("planner", "expected_reason"),
    [
        pytest.param(ProviderFailingPlanner(), "provider_unavailable", id="provider-failure"),
        pytest.param(RejectingPlanner(), "repair_rejected", id="rejected-repair"),
    ],
)
def test_unresolved_ordinary_span_after_failed_repair_is_audited_as_degraded(
    tmp_path: Path,
    planner: object,
    expected_reason: str,
) -> None:
    from docx import Document

    path = tmp_path / "semantic.docx"
    document = Document()
    document.add_paragraph("원문 유지")
    document.save(path)
    source = SourceFile(
        role=SourceRole.SEMANTIC_DOCUMENT,
        path=path,
        relative_path="semantic/semantic.docx",
        sha256=hashlib.sha256(path.read_bytes()).hexdigest(),
        size_bytes=path.stat().st_size,
        snapshot=path.read_bytes(),
    )

    parsed = DoclingParser(
        converter=FakeConverter(SimpleNamespace(iterate_items=lambda: iter(()))),
        structure_repair_planner=planner,  # type: ignore[arg-type]
    ).parse(source)

    assert parsed.markdown == "원문 유지\n"
    assert [item.excerpt for item in parsed.evidence] == ["원문 유지"]
    assert parsed.semantic_fidelity is not None
    assert parsed.semantic_fidelity.status == "WARN"
    assert parsed.semantic_fidelity.degraded_block_count == 1
    assert parsed.semantic_fidelity.degraded_blocks[0].reason == expected_reason
    assert parsed.semantic_fidelity.preserved_span_count == 1
    assert parsed.semantic_fidelity.source_text_coverage == 1.0
    assert parsed.semantic_fidelity.unmatched_span_count == 0
    assert parsed.semantic_fidelity.duplicated_span_count == 0


def test_residual_ordinary_span_after_accepted_repair_is_structure_unresolved(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    native = controlled_native_paragraphs(("FIRST", "SECOND"))
    first, _second = (span.span_id for span in native.spans)
    repair = RepairBlock(
        kind="paragraph",
        order=10,
        span_ids=[first],
        heading_level=None,
        list_kind=None,
        list_depth=None,
        row_count=None,
        column_count=None,
        cells=[],
        exclusion_kind=None,
        confidence=1.0,
    )

    parsed = parse_controlled(
        tmp_path,
        monkeypatch,
        native,
        planner=FixedRepairPlanner(
            RepairPlan(blocks=[repair]),
            (ParagraphBlock(order=10, span_ids=(first,)),),
        ),
    )

    assert parsed.markdown == "FIRST\n\nSECOND\n"
    assert [item.excerpt for item in parsed.evidence] == ["FIRST", "SECOND"]
    assert parsed.semantic_repair.outcome == "applied"
    assert parsed.semantic_fidelity.status == "WARN"
    assert parsed.semantic_fidelity.paragraph_count == 1
    assert parsed.semantic_fidelity.degraded_block_count == 1
    assert parsed.semantic_fidelity.degraded_blocks[0].reason == "structure_unresolved"
    assert parsed.semantic_fidelity.preserved_span_count == 2
    assert parsed.semantic_fidelity.source_text_coverage == 1.0
    assert parsed.semantic_fidelity.unmatched_span_count == 0
    assert parsed.semantic_fidelity.duplicated_span_count == 0


def test_full_page_ocr_preserves_paragraph_before_later_heading(
    tmp_path: Path,
) -> None:
    document = SimpleNamespace(
        pages={1: SimpleNamespace(size=SimpleNamespace(width=100.0, height=200.0))},
        iterate_items=lambda: iter(
            [
                (TextItem("FIRST", (0.05, 0.80, 0.20, 0.90)), 1),
                (SectionHeaderItem("SECOND", (0.05, 0.60, 0.20, 0.70)), 1),
            ]
        ),
    )

    parsed = DoclingParser(full_page_ocr_converter=FakeConverter(document)).parse(
        semantic_pdf_source(tmp_path)
    )

    assert parsed.markdown == "FIRST\n\n# SECOND\n"
    assert [item.excerpt for item in parsed.evidence] == ["FIRST", "SECOND"]
    assert parsed.semantic_fidelity is not None
    assert parsed.semantic_fidelity.status == "WARN"
    assert parsed.semantic_fidelity.heading_count == 1
    assert parsed.semantic_fidelity.paragraph_count == 1
    assert parsed.semantic_fidelity.degraded_block_count == 0
    assert parsed.semantic_fidelity.preserved_span_count == 2
    assert parsed.semantic_fidelity.source_text_coverage == 1.0
    assert parsed.semantic_fidelity.unmatched_span_count == 0
    assert parsed.semantic_fidelity.duplicated_span_count == 0


def test_unresolved_full_page_ocr_invokes_structure_repair(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    native = controlled_ocr_native(("Semantics 문서", "개인정보"))
    span_ids = tuple(span.span_id for span in native.spans)
    plan = RepairPlan(
        blocks=[
            RepairBlock(
                kind="paragraph",
                order=0,
                span_ids=list(span_ids),
                heading_level=None,
                list_kind=None,
                list_depth=None,
                row_count=None,
                column_count=None,
                cells=[],
                exclusion_kind=None,
                confidence=1.0,
            )
        ]
    )
    planner = FixedRepairPlanner(
        plan,
        (ParagraphBlock(order=0, span_ids=span_ids),),
    )

    parsed = parse_controlled(tmp_path, monkeypatch, native, planner=planner)

    assert parsed.markdown == "Semantics 문서개인정보\n"
    assert [item.excerpt for item in parsed.evidence] == ["Semantics 문서", "개인정보"]
    assert parsed.semantic_repair.outcome == "applied"
    assert parsed.semantic_repair.applied_orders == [0]
    assert parsed.semantic_fidelity.degraded_block_count == 0
    assert parsed.semantic_fidelity.source_text_coverage == 1.0


def test_semantic_structure_degraded_finding_includes_safe_repair_diagnostics(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    from ard_ossie.pipeline import _semantic_hard_findings

    native = controlled_ocr_native(("Semantics 문서", "개인정보"))
    parsed = parse_controlled(
        tmp_path,
        monkeypatch,
        native,
        planner=DetailedRejectingPlanner(),
    )

    findings = _semantic_hard_findings(parsed, require_visual_correction=False)

    specific = findings[0]
    assert specific.code == "SEMANTIC_REPAIR_ORDER_INVALID"
    assert "category=SEMANTIC_STRUCTURE_DEGRADED" in specific.message
    assert "extraction_mode=ocr" in specific.message
    assert "unresolved_spans=2" in specific.message
    assert "pages=1" in specific.message
    assert (
        "validation_codes=SEMANTIC_REPAIR_MISSING_SPAN,"
        "SEMANTIC_REPAIR_ORDER_INVALID"
    ) in specific.message
    assert "provider=fake" in specific.message
    assert "model=fake" in specific.message
    assert "attempts=2" in specific.message
    assert "Semantics 문서" not in specific.message
    assert any(item.code == "SEMANTIC_STRUCTURE_DEGRADED" for item in findings)


def test_full_page_ocr_converter_forces_ocr_for_every_pdf_page(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    import docling.datamodel.base_models as base_models
    import docling.datamodel.pipeline_options as pipeline_options
    import docling.document_converter as document_converter

    captured: dict[str, object] = {}

    class FakeEasyOcrOptions:
        def __init__(self, *, lang: list[str], force_full_page_ocr: bool) -> None:
            self.lang = lang
            self.force_full_page_ocr = force_full_page_ocr

    class FakePipelineOptions:
        def __init__(self, *, do_ocr: bool, ocr_options: object) -> None:
            self.do_ocr = do_ocr
            self.ocr_options = ocr_options

    class FakePdfFormatOption:
        def __init__(self, *, pipeline_options: object) -> None:
            self.pipeline_options = pipeline_options

    class CapturingConverter:
        def __init__(self, *, format_options: dict[object, object]) -> None:
            captured.update(format_options)

    monkeypatch.setattr(base_models, "InputFormat", SimpleNamespace(PDF="pdf"))
    monkeypatch.setattr(pipeline_options, "EasyOcrOptions", FakeEasyOcrOptions)
    monkeypatch.setattr(pipeline_options, "PdfPipelineOptions", FakePipelineOptions)
    monkeypatch.setattr(document_converter, "DocumentConverter", CapturingConverter)
    monkeypatch.setattr(document_converter, "PdfFormatOption", FakePdfFormatOption)

    _new_full_page_ocr_converter()

    options = captured["pdf"].pipeline_options
    assert options.do_ocr is True
    assert options.ocr_options.force_full_page_ocr is True
    assert options.ocr_options.lang == ["ko", "en"]


def test_docling_adapter_records_item_evidence_without_page_provenance(
    tmp_path: Path,
) -> None:
    path = tmp_path / "product.html"
    path.write_text("<html><body>사용자가 입력한 제품 목적</body></html>", encoding="utf-8")
    source = SourceFile(
        role=SourceRole.PRODUCT_HTML,
        path=path,
        relative_path="product-info/product.html",
        sha256=hashlib.sha256(path.read_bytes()).hexdigest(),
        size_bytes=path.stat().st_size,
        snapshot=path.read_bytes(),
    )

    parsed = DoclingParser(converter=FakeConverter(FakeHtmlDocument())).parse(source)

    assert parsed.evidence == [
        Evidence(
            source_hash=source.sha256,
            role=SourceRole.PRODUCT_HTML,
            locator={
                "document": "product-info/product.html",
                "item_index": 0,
                "level": 2,
            },
            excerpt="사용자가 입력한 제품 목적",
        )
    ]


def test_docling_adapter_excludes_ai_generated_label_and_adjacent_value_from_fact_evidence(
    tmp_path: Path,
) -> None:
    path = tmp_path / "product.html"
    path.write_text("<html><body>AI summary fixture</body></html>", encoding="utf-8")
    source = SourceFile(
        role=SourceRole.PRODUCT_HTML,
        path=path,
        relative_path="product-info/product.html",
        sha256=hashlib.sha256(path.read_bytes()).hexdigest(),
        size_bytes=path.stat().st_size,
        snapshot=path.read_bytes(),
    )

    parsed = DoclingParser(converter=FakeConverter(FakeAiGeneratedHtmlDocument())).parse(source)

    assert [item.excerpt for item in parsed.evidence] == ["사용자 설명"]
    assert [item.excerpt for item in parsed.excluded_product_fact_evidence] == [
        "(AI 자동생성) 데이터 요약",
        "자동 요약 값",
    ]
    assert "excluded_product_fact_evidence" not in parsed.model_dump(mode="json")


def test_real_docling_excludes_ai_generated_heading_and_direct_child_value(
    tmp_path: Path,
) -> None:
    path = tmp_path / "product.html"
    path.write_text(
        """<html><body>
        <h2>(AI 자동생성) 데이터 요약</h2>
        <p>자동 생성된 요약 값</p>
        <h2>사용자 설명</h2>
        <p>사용자가 작성한 설명 값</p>
        </body></html>""",
        encoding="utf-8",
    )
    source = SourceFile(
        role=SourceRole.PRODUCT_HTML,
        path=path,
        relative_path="product-info/product.html",
        sha256=hashlib.sha256(path.read_bytes()).hexdigest(),
        size_bytes=path.stat().st_size,
        snapshot=path.read_bytes(),
    )

    parsed = DoclingParser().parse(source)

    assert [item.excerpt for item in parsed.excluded_product_fact_evidence] == [
        "(AI 자동생성) 데이터 요약",
        "자동 생성된 요약 값",
    ]
    assert [item.excerpt for item in parsed.evidence] == [
        "사용자 설명",
        "사용자가 작성한 설명 값",
    ]


def test_real_docling_converts_html_and_docx_without_remote_service(tmp_path: Path) -> None:
    from docx import Document

    html_path = tmp_path / "product.html"
    html_path.write_text(
        "<html><body><h1>Sales Order</h1><p>Order analytics product.</p></body></html>",
        encoding="utf-8",
    )
    docx_path = tmp_path / "semantic.docx"
    document = Document()
    document.add_heading("Net Revenue", level=1)
    document.add_paragraph("Net revenue excludes tax.")
    document.save(docx_path)

    parser = DoclingParser()
    html = parser.parse(
        SourceFile(
            role=SourceRole.PRODUCT_HTML,
            path=html_path,
            relative_path="product-info/product.html",
            sha256=hashlib.sha256(html_path.read_bytes()).hexdigest(),
            size_bytes=html_path.stat().st_size,
            snapshot=html_path.read_bytes(),
        )
    )
    docx = parser.parse(
        SourceFile(
            role=SourceRole.SEMANTIC_DOCUMENT,
            path=docx_path,
            relative_path="semantic/semantic.docx",
            sha256=hashlib.sha256(docx_path.read_bytes()).hexdigest(),
            size_bytes=docx_path.stat().st_size,
            snapshot=docx_path.read_bytes(),
        )
    )

    assert "Sales Order" in html.markdown
    assert "Net Revenue" in docx.markdown
    assert "excludes tax" in docx.markdown
